from docx import Document


document = Document()


def header(text):
    document.add_heading(text, 0)


def paragrapher(text):
    document.add_paragraph(text).italic = True


def saver(name):
    document.save(name)